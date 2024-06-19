import pandas as pd
import pypdf as PyPDF2#PyPDF2
import docx
import glob
import re

# file folder location 
folder = r"C:\\Folder"

# read in all word/docx and pdf files
list_word_files = glob.glob(rf"{folder}/*.docx") + glob.glob(rf"{folder}/*.doc")
list_pdf_files = glob.glob(rf"{folder}/*.pdf")

############################### word documents ################################
# read word documents
word_documents_df = pd.DataFrame()
word_documents_df_sentences = pd.DataFrame()

word_documents_tables_df = pd.DataFrame()
word_documents_tables_sentences = pd.DataFrame()

for word_document in list_word_files:
    doc = docx.Document(word_document)
    paragraphs = doc.paragraphs
    tables = doc.tables
    
    # read paragraphs
    para_text = '\n'.join([p.text for p in paragraphs])
    #para_text_sentences = re.split(r"(?!\d)[.](?!\d?(?=[\w+]))", para_text) # split only on fullstops not between numbers
    para_text_sentences = re.split(r"(?<!\d)\.(?!\w|\d)", para_text) # split only on fullstop not between words or numbers
    word_doc_df = pd.DataFrame({'text':[para_text]})#re.split(r"(?!\d)[.](?!\d?(?=[\w+]))", para_text)})
    #word_doc_df['text'] = re.sub(r'\s+', ' ', word_doc_df['text']) # remove multiple whitespaces, replace with single whitespace
    word_doc_df['document'] = word_document
    word_documents_df = pd.concat([word_documents_df, word_doc_df], ignore_index=True)

    
    word_doc_df_sentences = pd.DataFrame({'text':para_text_sentences})
    word_doc_df_sentences['document'] = word_document
    word_documents_df_sentences = pd.concat([word_documents_df_sentences, word_doc_df_sentences], ignore_index=True)
    
    #for paragraph in paragraphs:
    #    para_text = '\n'.join([p.text for p in paragraphs])
    
    # read tables
    table_data, dictionary = [],{}
    table_data_sentences = []
    for count,table in enumerate(tables):
        keys = None
        for i,row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            
            if i == 0:
                keys = tuple(text)
                continue
            
            row_data = dict(zip(keys, text))
            row_data_sentences = ','.join('{}={!r}'.format(k, v) for k, v in row_data.items())
            table_data_sentences.append(row_data_sentences)
            table_data.append(row_data)
            
            dictionary[count] = row_data
            
    table_df = pd.DataFrame({'text':[str(table_data)]})
    table_df['document'] = word_document
    table_df_sentences = pd.DataFrame({'text':table_data_sentences, 'document':word_document})
    word_documents_tables_df = pd.concat([word_documents_tables_df,table_df], ignore_index=True)
    word_documents_tables_sentences = pd.concat([word_documents_tables_sentences,table_df_sentences], ignore_index=True)
    #word_documents_tables_df.to_csv(rf"{folder}\word_docs_tables_combined.csv", index=False)
    #word_documents_tables_sentences.to_csv(rf"{folder}\word_docs_tables_sentences.csv", index=False)
    
    # join table text
    table_text = '\n'.join([p for p in word_documents_tables_sentences['text']])
    word_tbls = pd.DataFrame({'text':table_text,'document':[word_document]})


##################################### pdf documents ###########################
# read pdf documents
pdf_documents_df = pd.DataFrame()
pdf_documents_df_sentences = pd.DataFrame()

for pdf_document in list_pdf_files:
    pdf_files = PyPDF2.PdfReader(pdf_document)
    pdf_text = ""

    for page in pdf_files.pages:
        pdf_text += page.extract_text() + "\n"
    pdf_df = pd.DataFrame({'text':pdf_text,'document':[pdf_document]})
    pdf_df_sentences = pd.DataFrame({'text':re.split(r"(?!\d)[.](?!\d?(?=[\w+]))",pdf_text), 'document':pdf_document})
    
    pdf_documents_df = pd.concat([pdf_documents_df,pdf_df], ignore_index=True)
    #pdf_documents_df.to_csv(rf"{folder}\pdf_docs_combined.csv", index=False)
    
    pdf_documents_df_sentences = pd.concat([pdf_documents_df_sentences,pdf_df_sentences], ignore_index=True)

# write csv with all of these
df = pd.concat([word_documents_df,word_tbls], ignore_index=True)
df = pd.concat([df,pdf_documents_df], ignore_index=True)
df.to_csv(rf"{folder}\dfs_combined.csv", index=False)

# df sentences
df_sentences = pd.concat([word_documents_df_sentences, word_documents_tables_sentences], ignore_index = True)
df_sentences = pd.concat([df_sentences, pdf_documents_df_sentences], ignore_index = True)
df_sentences.to_csv(rf"{folder}\dfs_combined_sentences.csv", index=False)

################################  text chunking ###############################
# split long strings of text into chunks of text
# excel csv has character read limit of 32,767 per cell
# make edited df which keeps to this limt or less per row

edited_df = df

# if string length > 32,767, split data on a fullstop (not between numbers) around character 15,000
# repeat process as many times as needed to get all entries at or below limit
# https://stackoverflow.com/questions/32122022/split-a-string-into-pieces-of-max-length-x-split-only-at-spaces
import textwrap

n = 10000

######################## textwrap.wrap function ###############################
new_text = []
temp_df = pd.DataFrame()
for each_line in range(0, len(edited_df['text']),1):
    document_name = edited_df['document'][each_line]
    each_line1 = edited_df['text'][each_line]
    lines = textwrap.wrap(each_line1, n, break_long_words=False)
    
    # make df
    #new_text.append(lines)
    hold_df = pd.DataFrame({'text':lines})
    hold_df['document'] = document_name
    
    temp_df = pd.concat([temp_df,hold_df], ignore_index=True)


################# custom implementation #######################################
new_text = []
temp_df = pd.DataFrame()
for each_line in range(0, len(edited_df['text']),1):
    document_name = edited_df['document'][each_line]
    each_line1 = edited_df['text'][each_line]
    lines = []
    docer = []
    if len(each_line1) > n:
        words = iter(each_line1.split())
        lines, current = [], next(words)
        docer = []
        for word in words:
            if len(current) + 1 + len(word) > n:
                lines.append(current)
                docer.append(document_name)
                current = word
                #docer = document_name
            else:
                current += " " + word
                #document_name += document_name
        lines.append(current)
        docer.append(document_name)
            
    if len(each_line1) <= n:
        lines.append(each_line1)
        docer.append(document_name)
        
        
    # make df
    #new_text.append(lines)
    hold_df = pd.DataFrame({'text':lines})
    hold_df['document'] = document_name
    
    temp_df = pd.concat([temp_df,hold_df], ignore_index=True)
    
xl_csv = temp_df

#xl_csv = pd.DataFrame({'text':lines,'document':docer})
xl_csv.to_csv(rf"{folder}\dfs_combined_xl.csv", index=False)    

###############################################################################

# read outlook msg and attachment
import extract_msg

file_types = ["xls", "xlsx", "docx", "pdf"]

msg = extract_msg.openMsg("Msg.msg")
sender = msg.sender
subject = msg.subject
body = msg.body
time_received = msg.receivedTime # datetime
attachment_filenames = []
attachment_content = []

for att in msg.attachments:
    att.save(customPath=rf"{folder}") # save a copy of the attachment
    
    # open the attachment and get the attachment extension
    att_ext = re.findall("(?<=[.]).*",att.name)
    att_ext = att_ext[0]
    
    if att_ext not in file_types:
        valueError(f"Error - file {att_ext} not suitable for use")
    else:
        # read attachment
        if att_ext == "csv":
            pd.read_csv(att)
            
        elif att_ext == "xlsx" or att_ext == "xls":
            pd.read_excel(att)
            
        elif att_ext == "docx" or att_ext == "doc":
            doc = docx.Document(att)
            paragraphs = doc.paragraphs
            tables = doc.tables
            
            # read paragraphs
            para_text = '\n'.join([p.text for p in paragraphs])
            #para_text_sentences = re.split(r"(?!\d)[.](?!\d?(?=[\w+]))", para_text) # split only on fullstops not between numbers
            para_text_sentences = re.split(r"(?<!\d)\.(?!\w|\d)", para_text) # split only on fullstop not between words or numbers
            word_doc_df = pd.DataFrame({'text':[para_text]})#re.split(r"(?!\d)[.](?!\d?(?=[\w+]))", para_text)})
            #word_doc_df['text'] = re.sub(r'\s+', ' ', word_doc_df['text']) # remove multiple whitespaces, replace with single whitespace
            word_doc_df['document'] = word_document
            word_documents_df = pd.concat([word_documents_df, word_doc_df], ignore_index=True)

            
            word_doc_df_sentences = pd.DataFrame({'text':para_text_sentences})
            word_doc_df_sentences['document'] = word_document
            word_documents_df_sentences = pd.concat([word_documents_df_sentences, word_doc_df_sentences], ignore_index=True)
            
            #for paragraph in paragraphs:
            #    para_text = '\n'.join([p.text for p in paragraphs])
            
            # read tables
            table_data, dictionary = [],{}
            table_data_sentences = []
            for count,table in enumerate(tables):
                keys = None
                for i,row in enumerate(table.rows):
                    text = (cell.text for cell in row.cells)
                    
                    if i == 0:
                        keys = tuple(text)
                        continue
                    
                    row_data = dict(zip(keys, text))
                    row_data_sentences = ','.join('{}={!r}'.format(k, v) for k, v in row_data.items())
                    table_data_sentences.append(row_data_sentences)
                    table_data.append(row_data)
                    
                    dictionary[count] = row_data
                    
            table_df = pd.DataFrame({'text':[str(table_data)]})
            table_df['document'] = word_document
            table_df_sentences = pd.DataFrame({'text':table_data_sentences, 'document':word_document})
            word_documents_tables_df = pd.concat([word_documents_tables_df,table_df], ignore_index=True)
            word_documents_tables_sentences = pd.concat([word_documents_tables_sentences,table_df_sentences], ignore_index=True)
            #word_documents_tables_df.to_csv(rf"{folder}\word_docs_tables_combined.csv", index=False)
            #word_documents_tables_sentences.to_csv(rf"{folder}\word_docs_tables_sentences.csv", index=False)
            
            # join table text
            table_text = '\n'.join([p for p in word_documents_tables_sentences['text']])
            word_tbls = pd.DataFrame({'text':table_text,'document':[word_document]})
            
        elif att_ext == "pdf":
          pdf_files = PyPDF2.PdfReader(att)
          pdf_text = ""
          for page in pdf_files.pages:
              pdf_text += page.extract_text() + "\n"  
        
            
# capture msg content and put in table
msg_df = pd.DataFrame({"sender":[sender], "time_sent":[str(time_received)], "subject":[subject], "content":[body]})
msg_df.to_csv(rf"{output_folder}\file_msg.csv", index=False)